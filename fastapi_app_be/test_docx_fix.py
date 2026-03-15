import io
import sys
import os

# Add the project directory to sys.path for imports
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

from docx import Document
from services.docx_service import _sanitize_text, _format_projects, _delete_paragraph
from models.resume import ResumeData, ProjectExperience, ProjectDetail

def test_sanitization():
    print("Testing sanitization...")
    bad_text = "Line 1\x0cLine 2\x00Line 3\x07"
    sanitized = _sanitize_text(bad_text)
    print(f"Original length: {len(bad_text)}")
    print(f"Sanitized length: {len(sanitized)}")
    assert "\x0c" not in sanitized
    assert "\x00" not in sanitized
    assert "\x07" not in sanitized
    assert "Line 1Line 2Line 3" in sanitized
    print("Sanitization test passed.")

def test_empty_table_prevention():
    print("\nTesting empty table prevention...")
    doc = Document()
    p = doc.add_paragraph("PROJECT_EXPERIENCE")
    
    # Case 1: None/Empty list
    _format_projects(p, None)
    # Check if a table was added to the document
    assert len(doc.tables) == 0
    print("Case 1 (None) passed: No table created.")
    
    # Case 2: List with invalid/empty projects
    projects = [ProjectExperience(description="", projectDetails=[], responsibilities=[])]
    _format_projects(p, projects)
    assert len(doc.tables) == 0
    print("Case 2 (Empty project) passed: No table created.")
    
    # Case 3: List with valid project
    projects = [ProjectExperience(description="Valid Project", projectDetails=[], responsibilities=[])]
    _format_projects(p, projects)
    assert len(doc.tables) == 1
    print("Case 3 (Valid project) passed: 1 table created.")

def test_deletion_logic():
    print("\nTesting deletion logic...")
    doc = Document()
    doc.add_paragraph("Title")
    p = doc.add_paragraph("PLACEHOLDER")
    
    # Before deletion
    assert len(doc.paragraphs) == 2
    
    # Delete p and title (at p_idx - 1)
    from services.docx_service import _inject_section
    _inject_section(p, None, lambda x, y: None)
    
    # After deletion, both should be gone
    assert len(doc.paragraphs) == 0
    print("Deletion logic test passed.")

if __name__ == "__main__":
    try:
        test_sanitization()
        test_empty_table_prevention()
        test_deletion_logic()
        print("\nALL VERIFICATION TESTS PASSED SUCCESSFULLY.")
    except Exception as e:
        print(f"\nVERIFICATION FAILED: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
