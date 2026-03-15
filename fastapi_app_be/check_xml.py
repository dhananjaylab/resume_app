import io
from docx import Document
import xml.dom.minidom

def check_empty_table_xml():
    doc = Document()
    p = doc.add_paragraph("Para")
    table = doc.add_table(rows=0, cols=2)
    p._p.addprevious(table._tbl)
    
    # Save and look at XML
    buf = io.BytesIO()
    doc.save(buf)
    
    # We can't easily read the zip here without more boilerplate, 
    # but let's see what table._tbl looks like
    from docx.oxml.ns import qn
    print("Table XML element tag:", table._tbl.tag)
    print("Number of rows in XML:", len(table._tbl.findall(qn('w:tr'))))

if __name__ == "__main__":
    check_empty_table_xml()
