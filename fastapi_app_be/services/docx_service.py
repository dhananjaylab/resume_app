"""
DOCX service for generating and downloading resume documents.
Creates formatted DOCX files from ResumeData using a template.
Updated for new field names (candidateName, from_date, etc.)
"""
import io
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models.resume import ResumeData
from exceptions.custom_exceptions import FileProcessingException


async def download_resume(resume_data: ResumeData, template_path: str = None) -> tuple[bytes, str]:
    """
    Generate a DOCX file from ResumeData and return as bytes.
    
    Args:
        resume_data: ResumeData object containing resume information
        template_path: Path to DOCX template (optional)
        
    Returns:
        Tuple of (file_bytes, filename)
        
    Raises:
        FileProcessingException: If DOCX generation fails
    """
    try:
        # Use provided template or default
        if template_path is None:
            template_path = "templates/Maveric_Template.docx"
        
        # Check if template exists
        if not os.path.exists(template_path):
            # If template doesn't exist, create document from scratch
            doc = Document()
        else:
            # Load template
            doc = Document(template_path)
        
        # Populate document with resume data
        _populate_document(doc, resume_data)
        
        # Generate filename
        filename = _generate_filename(resume_data)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue(), filename
        
    except Exception as e:
        raise FileProcessingException(f"Error generating DOCX file: {str(e)}")


def _populate_document(doc: Document, resume_data: ResumeData) -> None:
    """
    Populate document with resume data.
    Finds placeholders like SUMMARY, EXPERIENCE, etc. and replaces them.
    
    Args:
        doc: Document object to populate
        resume_data: ResumeData object with information
    """
    # Replace Header Name & Position in the exact first table
    if resume_data.headers:
        _replace_header_table(doc, resume_data.headers)
    
    # Process each placeholder section
    # Note: We iterate over a snapshot of paragraphs to avoid mutation issues
    for p in list(doc.paragraphs):
        text = p.text.strip()
        
        if "SUMMARY" in text:
            _inject_section(p, resume_data.professionalSummary, _format_summary)
        elif "EXPERIENCE" in text:
            _inject_section(p, resume_data.professionalExperience, _format_experience)
        elif "AWARDS" in text:
            _inject_section(p, resume_data.awards, _format_list)
        elif "CERTIFICATION" in text:
            _inject_section(p, resume_data.certifications, _format_list)
        elif "EDUCATION" in text:
            _inject_section(p, resume_data.education, _format_list)
        elif "CREDITS" in text:
            _inject_section(p, resume_data.credits, _format_credits)
        elif "PROJECT_EXPERIENCE" in text:
            _inject_section(p, resume_data.projectExperience, _format_projects)


def _replace_header_table(doc: Document, headers) -> None:
    """Find the table in the header containing NAME and POSITION and replace them."""
    for section in doc.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Update cell text if matches exactly or contains placeholder
                    if "NAME" in cell.text and headers.candidateName:
                        _replace_cell_text(cell, "NAME", headers.candidateName, bold=True, size=Pt(16))
                    if "POSITION" in cell.text and headers.candidatePosition:
                        _replace_cell_text(cell, "POSITION", headers.candidatePosition, bold=False, size=Pt(12))


def _replace_cell_text(cell, placeholder: str, new_text: str, bold: bool = False, size=None):
    """Safely replace text in a table cell while keeping alignment."""
    cell.text = cell.text.replace(placeholder, new_text)
    for p in cell.paragraphs:
        for r in p.runs:
            if new_text in r.text:
                r.font.bold = bold
                if size:
                    r.font.size = size


def _inject_section(placeholder_paragraph, data, formatter_fn):
    """
    Inject section data before the placeholder paragraph.
    If no data, delete the placeholder and its preceding title.
    """
    if data:
        formatter_fn(placeholder_paragraph, data)
        _delete_paragraph(placeholder_paragraph)
    else:
        # If no data, find the previous paragraph (which is usually the title) and delete it too
        prev_p = placeholder_paragraph.insert_paragraph_before("")
        p_idx = placeholder_paragraph._element.getparent().index(placeholder_paragraph._element)
        parent = placeholder_paragraph._element.getparent()
        
        # We inserted prev_p exactly before, so prev_p's index is p_idx.
        # The real title paragraph is at p_idx - 1
        if p_idx > 0:
            title_p_element = parent[p_idx - 1]
            parent.remove(title_p_element)
            
        _delete_paragraph(placeholder_paragraph)
        _delete_paragraph(prev_p)


def _delete_paragraph(paragraph):
    """Safely remove a paragraph from the Document."""
    p_element = paragraph._element
    p_parent = p_element.getparent()
    if p_parent is not None:
        p_parent.remove(p_element)


# --- FORMATTERS ---

def _format_summary(p, summary: str):
    if not summary: return
    new_p = p.insert_paragraph_before(summary)
    new_p.space_after = Pt(12)


def _format_experience(p, experiences: list):
    for exp in experiences:
        if exp and isinstance(exp, str):
            new_p = p.insert_paragraph_before(f"• {exp}")
            new_p.space_after = Pt(6)


def _format_list(p, items: list):
    for item in items:
        if item and isinstance(item, str):
            new_p = p.insert_paragraph_before(f"• {item}")
            new_p.space_after = Pt(6)


def _format_credits(p, credits_list: list):
    for credit in credits_list:
        if not credit.category:
            continue
        
        cat_p = p.insert_paragraph_before()
        cat_run = cat_p.add_run(f"{credit.category}:")
        cat_run.font.bold = True
        cat_p.space_before = Pt(6)
        
        if credit.items:
            items_text = ", ".join(credit.items)
            item_p = p.insert_paragraph_before(items_text)
            item_p.space_after = Pt(6)


def _format_projects(p, projects: list):
    if not projects:
        return
        
    doc = p._parent
    table = doc.add_table(rows=0, cols=2)
    p._p.addprevious(table._tbl)
    
    for proj in projects:
        if not proj or not (proj.description or proj.projectDetails or proj.responsibilities):
            continue
            
        row = table.add_row()
        cell_0 = row.cells[0]
        cell_1 = row.cells[1]
        
        # Column 0: Details
        cell_0.width = Inches(2.0)
        c0_p = cell_0.paragraphs[0]
        
        if proj.projectDetails:
            first0 = True
            for detail in proj.projectDetails:
                if detail.key and detail.value:
                    if not first0:
                        c0_p = cell_0.add_paragraph()
                    first0 = False
                    
                    k = detail.key.strip()
                    k = k[0].upper() + k[1:] if k else ""
                    v = detail.value.strip()
                    v = v[0].upper() + v[1:] if v else ""
                    
                    r_k = c0_p.add_run(f"{k}: ")
                    r_k.font.bold = True
                    c0_p.add_run(v)
        
        # Column 1: Description and Responsibilities
        cell_1.width = Inches(4.5)
        c1_p = cell_1.paragraphs[0]
        
        c1_first = True
        if proj.description:
            r_d = c1_p.add_run("Description:\n")
            r_d.font.bold = True
            c1_p.add_run(proj.description)
            c1_first = False
            
        if proj.responsibilities and len(proj.responsibilities) > 0:
            if not c1_first:
                c1_p = cell_1.add_paragraph()
            r_r = c1_p.add_run("Responsibilities:")
            r_r.font.bold = True
            
            for index, resp in enumerate(proj.responsibilities):
                if resp:
                    resp_p = cell_1.add_paragraph(f"• {resp}")
                    if index == len(proj.responsibilities) - 1:
                        resp_p.space_after = Pt(12)
        
        # Spacing Row
        spacer = table.add_row()


def _generate_filename(resume_data: ResumeData) -> str:
    """
    Generate a filename for the resume based on candidate name.
    
    Args:
        resume_data: ResumeData object with resume information
        
    Returns:
        Filename string for the DOCX file
    """
    if resume_data.headers and resume_data.headers.candidateName:
        name = resume_data.headers.candidateName.replace(" ", "_")
        return f"{name}_Resume.docx"
    else:
        return "Resume.docx"
