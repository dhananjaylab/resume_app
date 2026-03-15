import io
from docx import Document
from docx.shared import Pt

def test_empty_table():
    print("Testing 0-row table...")
    doc = Document()
    p = doc.add_paragraph("Testing table insertion")
    table = doc.add_table(rows=0, cols=2)
    p._p.addprevious(table._tbl)
    
    buf = io.BytesIO()
    doc.save(buf)
    print("0-row table saved successfully.")

def test_control_characters():
    print("Testing control characters...")
    doc = Document()
    # \x0c is Form Feed, often found in resumes (page breaks)
    # \x00 - \x08, \x0b, \x0c, \x0e - \x1f are invalid in XML 1.0
    bad_text = "Hello\x0cWorld"
    try:
        doc.add_paragraph(bad_text)
        buf = io.BytesIO()
        doc.save(buf)
        print("Control characters (\\x0c) saved successfully (python-docx didn't complain).")
    except Exception as e:
        print(f"python-docx failed to save control characters: {e}")

if __name__ == "__main__":
    test_empty_table()
    test_control_characters()
